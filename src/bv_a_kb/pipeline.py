from __future__ import annotations

import json
import mimetypes
import re
import shutil
import subprocess
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse

import cv2
import gdown
import numpy as np
import openpyxl
import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR

from .config import (
    ALLOWED_PATTERNS,
    BASE_URL,
    CANONICAL_DOC_DIR,
    CATALOG_DIR,
    DATA_DIR,
    DEFAULT_TIMEOUT,
    DIRECT_FILE_EXTENSIONS,
    EXCLUDED_PATTERNS,
    KATANA_DEPTH_BY_MODE,
    KATANA_DURATION_BY_MODE,
    KNOWLEDGE_BASE_DIR,
    PROCESS_OCR_PHRASES,
    RAW_ASSET_DIR,
    RAW_EMBED_DIR,
    RAW_HTML_DIR,
    SEED_URLS,
    SMOKE_TARGET_URLS,
    STATIC_RESOURCE_EXTENSIONS,
    TOPIC_GROUP_BY_URL,
    USER_AGENT,
)
from .models import EmbeddedResource, ExtractedTable, KnowledgeDocument, PipelineReport, SourceRecord
from .utils import (
    absolutize,
    ascii_fold,
    drop_boilerplate_lines,
    ensure_dirs,
    make_id,
    normalize_text,
    slug_from_url,
    write_json,
    write_jsonl,
)


class BVAKnowledgeBasePipeline:
    def __init__(self) -> None:
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": USER_AGENT})
        self.ocr_engine = RapidOCR()
        ensure_dirs(
            [
                DATA_DIR,
                CATALOG_DIR,
                RAW_HTML_DIR,
                RAW_ASSET_DIR,
                RAW_EMBED_DIR,
                KNOWLEDGE_BASE_DIR,
                CANONICAL_DOC_DIR,
            ]
        )

    def classify_topic_group(self, url: str) -> str:
        lowered = url.lower()
        if "/his/contact/" in lowered:
            return "department"
        for prefix, topic_group in TOPIC_GROUP_BY_URL.items():
            if prefix in url:
                return topic_group
        return "generic"

    def classify_page_type(self, url: str) -> str:
        lowered = url.lower()
        if "/his/contact/" in lowered:
            return "department_contact"
        if "/faqs/faq/" in lowered:
            return "faq_detail"
        if lowered.rstrip("/") == f"{BASE_URL}/faqs" or "/faqs/index/" in lowered:
            return "faq_listing"
        if "/his/department/" in lowered:
            return "department"
        if "/articles/category/" in lowered:
            return "article_listing"
        if "/article/" in lowered:
            return "article"
        if lowered.rstrip("/") == f"{BASE_URL}/contact":
            return "contact"
        if "/page/" in lowered:
            return "page"
        if "/docs/category/" in lowered:
            return "document_listing"
        if "/vaccines" in lowered:
            return "service_page"
        if "/tokhaiyte" in lowered or "/his/book_appointment" in lowered:
            return "form_page"
        return "generic"

    def should_keep_url(self, url: str) -> bool:
        normalized = absolutize(url)
        lowered = normalized.lower()
        if not lowered.startswith(BASE_URL):
            return False
        if any(lowered.endswith(ext) for ext in STATIC_RESOURCE_EXTENSIONS):
            return False
        if any(pattern in lowered for pattern in EXCLUDED_PATTERNS):
            return False
        if lowered.rstrip("/") == f"{BASE_URL}/contact":
            return True
        if any(lowered.endswith(ext) for ext in DIRECT_FILE_EXTENSIONS):
            return True
        if "/article/" in lowered:
            keep_article = any(
                token in lowered
                for token in (
                    "bhyt",
                    "bao-hiem",
                    "huong-dan",
                    "dich-vu",
                    "kham",
                    "suc-khoe",
                    "vien-phi",
                    "tiem-chung",
                    "quy-trinh",
                    "thu-tuc",
                )
            )
            if not keep_article:
                return False
        return any(pattern in lowered for pattern in ALLOWED_PATTERNS)

    def _build_source_record(
        self,
        source_url: str,
        discovered_from: str | None = None,
        source_tag: str | None = None,
        source_attribute: str | None = None,
        status_code: int | None = None,
        content_length: int | None = None,
        crawl_status: str | None = None,
    ) -> SourceRecord:
        return SourceRecord(
            source_url=source_url,
            page_type=self.classify_page_type(source_url),
            topic_group=self.classify_topic_group(source_url),
            discovered_from=discovered_from,
            source_tag=source_tag,
            source_attribute=source_attribute,
            status_code=status_code,
            content_length=content_length,
            crawl_status=crawl_status or ("ok" if (status_code or 0) < 400 else "pending"),
        )

    def _merge_source_record(self, source_map: dict[str, SourceRecord], record: SourceRecord) -> None:
        existing = source_map.get(record.source_url)
        if existing is None:
            source_map[record.source_url] = record
            return
        if (existing.status_code or 0) >= 400 > (record.status_code or 0):
            source_map[record.source_url] = record
            return
        if existing.discovered_from is None and record.discovered_from:
            existing.discovered_from = record.discovered_from
            existing.source_tag = existing.source_tag or record.source_tag
            existing.source_attribute = existing.source_attribute or record.source_attribute

    def _iter_relevant_links(self, html: str, page_url: str) -> Iterable[tuple[str, str, str]]:
        soup = BeautifulSoup(html, "lxml")
        for tag_name, attr_name in (("a", "href"), ("area", "href")):
            for tag in soup.find_all(tag_name, attrs={attr_name: True}):
                href = tag.get(attr_name, "").strip()
                if not href:
                    continue
                yield tag_name, attr_name, absolutize(urljoin(page_url, href))

    def _supplement_discovery(self, source_map: dict[str, SourceRecord], max_rounds: int = 3) -> int:
        added = 0
        visited: set[str] = set()
        for _ in range(max_rounds):
            frontier = [
                source_url
                for source_url, record in sorted(source_map.items())
                if source_url not in visited
                and record.crawl_status == "ok"
                and not any(source_url.lower().endswith(ext) for ext in DIRECT_FILE_EXTENSIONS)
            ]
            if not frontier:
                break
            round_added = 0
            for page_url in frontier:
                visited.add(page_url)
                try:
                    html = self.fetch_html(page_url)
                except requests.RequestException:
                    continue
                for tag_name, attr_name, discovered_url in self._iter_relevant_links(html, page_url):
                    if not self.should_keep_url(discovered_url):
                        continue
                    if discovered_url in source_map:
                        continue
                    self._merge_source_record(
                        source_map,
                        self._build_source_record(
                            discovered_url,
                            discovered_from=page_url,
                            source_tag=tag_name,
                            source_attribute=attr_name,
                            crawl_status="ok",
                        ),
                    )
                    round_added += 1
            added += round_added
            if round_added == 0:
                break
        return added

    def _run_katana(self, seeds: list[str], depth: int, crawl_duration: str) -> list[dict]:
        if shutil.which("katana") is None:
            raise RuntimeError("katana was not found in PATH")

        seed_file = CATALOG_DIR / "katana_seeds.txt"
        seed_file.write_text("\n".join(seeds), encoding="utf-8")

        cmd = [
            "katana",
            "-list",
            str(seed_file),
            "-d",
            str(depth),
            "-ct",
            crawl_duration,
            "-j",
            "-silent",
            "-nc",
            "-jc",
            "-iqp",
            "-fsu",
            "-kf",
            "all",
            "-fs",
            "rdn",
            "-timeout",
            "15",
            "-retry",
            "2",
            "-mrs",
            str(8 * 1024 * 1024),
            "-or",
            "-ob",
            "-eof",
            "headers",
        ]
        completed = subprocess.run(
            cmd,
            cwd=str(DATA_DIR.parent),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if completed.returncode not in {0, 1}:
            raise RuntimeError(f"katana failed with exit code {completed.returncode}: {completed.stderr.strip()}")

        output_file = CATALOG_DIR / "katana_discovery.jsonl"
        output_file.write_text(completed.stdout, encoding="utf-8")

        events: list[dict] = []
        for line in completed.stdout.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                events.append(json.loads(line))
            except json.JSONDecodeError:
                continue
        return events

    def discover_sources(self, mode: str, depth: int, crawl_duration: str) -> list[SourceRecord]:
        events = self._run_katana(SEED_URLS, depth=depth, crawl_duration=crawl_duration)
        source_map: dict[str, SourceRecord] = {}
        kept_events = 0

        for event in events:
            request = event.get("request") or {}
            response = event.get("response") or {}
            endpoint = request.get("endpoint")
            if not endpoint:
                continue
            endpoint = absolutize(endpoint)
            if not self.should_keep_url(endpoint):
                continue
            kept_events += 1
            record = self._build_source_record(
                endpoint,
                discovered_from=request.get("source"),
                source_tag=request.get("tag"),
                source_attribute=request.get("attribute"),
                status_code=response.get("status_code"),
                content_length=response.get("content_length"),
                crawl_status="ok" if (response.get("status_code") or 0) < 400 else "error",
            )
            self._merge_source_record(source_map, record)

        supplemental_sources = self._supplement_discovery(source_map)
        records = sorted(source_map.values(), key=lambda item: item.source_url)

        write_json(CATALOG_DIR / "seeds.json", SEED_URLS)
        write_json(
            CATALOG_DIR / "crawl_report.json",
            {
                "mode": mode,
                "katana_events": len(events),
                "kept_events": kept_events,
                "supplemental_sources": supplemental_sources,
                "unique_sources": len(records),
                "depth": depth,
                "crawl_duration": crawl_duration,
            },
        )
        write_jsonl(CATALOG_DIR / "discovered_sources.jsonl", [record.model_dump() for record in records])
        return records

    def fetch_html(self, url: str) -> str:
        response = self.session.get(url, timeout=DEFAULT_TIMEOUT)
        response.raise_for_status()
        response.encoding = response.apparent_encoding or response.encoding or "utf-8"
        return response.text

    def select_main_container(self, soup: BeautifulSoup, page_type: str):
        selectors = {
            "contact": ["#content .col-md-9", ".main-content .container .row .col-md-9"],
            "department_contact": ["#content .col-md-9", ".main-content .container .row .col-md-9"],
            "page": ["#article-item", ".article-content", ".main-content .col-md-9"],
            "article": [".article-content", ".entry-content", ".main-content .col-md-9"],
            "article_listing": [".list-dashed", ".main-content .col-md-9"],
            "department": [".blog-posts", ".blog-pull-right", ".main-content .col-md-8"],
            "faq_listing": [".question-list", ".main-content .col-md-9"],
            "faq_detail": [".post.clearfix", ".entry-content", ".main-content .col-md-9"],
            "document_listing": [".main-content .col-md-9", "#content .col-md-9"],
            "service_page": [".main-content .col-md-9", ".article-content"],
            "form_page": [".main-content", "body"],
            "generic": ["#content", ".main-content", "body"],
        }
        for selector in selectors.get(page_type, selectors["generic"]):
            node = soup.select_one(selector)
            if node is not None:
                return node
        return soup.body or soup

    def _extract_breadcrumbs(self, soup: BeautifulSoup) -> list[str]:
        crumbs: list[str] = []
        for node in soup.select(".breadcrumb li, .breadcrumb li a"):
            value = normalize_text(node.get_text(" ", strip=True))
            if value and value not in crumbs:
                crumbs.append(value)
        return crumbs

    def _extract_text(self, container) -> str:
        removable = (
            "script, style, nav, header, footer, .widget, .panel, .breadcrumb, form, "
            ".addthis_inline_share_toolbox, .search-form-wrapper, .sidebar, .post-right"
        )
        for tag in container.select(removable):
            tag.decompose()
        text = container.get_text("\n", strip=True)
        return drop_boilerplate_lines(normalize_text(text))

    def _extract_structured_faq(self, container) -> tuple[str, str]:
        question = ""
        answer = ""
        for block in container.find_all("div", recursive=False):
            heading = block.find("h4")
            if heading is None:
                continue
            label = ascii_fold(normalize_text(heading.get_text(" ", strip=True)))
            cloned = BeautifulSoup(str(block), "lxml")
            cloned_heading = cloned.find("h4")
            if cloned_heading is not None:
                cloned_heading.decompose()
            content = normalize_text(cloned.get_text("", strip=False))
            if "cau hoi" in label and content:
                question = content
            elif "tra loi" in label and content:
                answer = content
        combined: list[str] = []
        if question:
            combined.append(f"Cau hoi:\n{question}")
        if answer:
            combined.append(f"Tra loi:\n{answer}")
        return "\n".join(combined).strip(), question[:120] if question else "FAQ"

    def _candidate_asset_urls(self, src: str) -> list[str]:
        candidates = [src]
        parsed = urlparse(src)
        beta_host = "beta.benhvienathainguyen.com.vn"
        main_host = urlparse(BASE_URL).netloc
        if parsed.netloc == beta_host:
            for scheme, host in (("https", main_host), ("http", main_host), ("https", beta_host)):
                candidate = parsed._replace(scheme=scheme, netloc=host).geturl()
                if candidate not in candidates:
                    candidates.append(candidate)
        return candidates

    def _extract_images(self, container, doc_slug: str) -> list[EmbeddedResource]:
        assets: list[EmbeddedResource] = []
        for index, img in enumerate(container.select("img[src]"), start=1):
            src = absolutize(img.get("src", ""))
            if not src.startswith("http"):
                continue
            if "/uploads/" not in src or ".thumbs/" in src:
                continue
            suffix = Path(urlparse(src).path).suffix or ".bin"
            local_path = RAW_ASSET_DIR / f"{doc_slug}-img-{index}{suffix}"
            status = "skipped"
            note = None
            resolved_url = src
            last_error = None
            for candidate_url in self._candidate_asset_urls(src):
                try:
                    response = self.session.get(candidate_url, timeout=DEFAULT_TIMEOUT)
                    response.raise_for_status()
                    local_path.write_bytes(response.content)
                    status = "downloaded"
                    resolved_url = candidate_url
                    if candidate_url != src:
                        note = "rewritten_from_beta_host"
                    break
                except requests.RequestException as exc:
                    last_error = str(exc)
            if status != "downloaded":
                note = last_error
            assets.append(
                EmbeddedResource(
                    kind="image",
                    source_url=src,
                    resolved_url=resolved_url,
                    local_path=str(local_path) if local_path.exists() else None,
                    mime_type=mimetypes.guess_type(str(local_path))[0],
                    status=status,
                    note=note,
                )
            )
        return assets

    def _download_embedded_resource(self, src: str, doc_slug: str, index: int, kind: str) -> EmbeddedResource:
        note = None
        local_path: Path | None = None
        status = "discovered"
        resolved_url = src
        lowered = src.lower()

        try:
            if "google.com/file/d/" in lowered:
                match = re.search(r"/file/d/([^/]+)/", src)
                if match:
                    file_id = match.group(1)
                    target_path = RAW_EMBED_DIR / f"{doc_slug}-{kind}-{index}"
                    gdown.download(f"https://drive.google.com/uc?id={file_id}", str(target_path), quiet=True, fuzzy=True)
                    local_path = target_path
                    status = "downloaded"
            elif "docs.google.com/document/d/" in lowered:
                match = re.search(r"/document/d/([^/]+)/", src)
                if match:
                    doc_id = match.group(1)
                    target_path = RAW_EMBED_DIR / f"{doc_slug}-{kind}-{index}.txt"
                    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
                    response = self.session.get(export_url, timeout=DEFAULT_TIMEOUT)
                    response.raise_for_status()
                    target_path.write_bytes(response.content)
                    local_path = target_path
                    resolved_url = export_url
                    status = "downloaded"
            elif "docs.google.com/spreadsheets/d/" in lowered:
                match = re.search(r"/spreadsheets/d/([^/]+)/", src)
                if match:
                    sheet_id = match.group(1)
                    target_path = RAW_EMBED_DIR / f"{doc_slug}-{kind}-{index}.xlsx"
                    export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx"
                    response = self.session.get(export_url, timeout=DEFAULT_TIMEOUT)
                    if response.status_code == 200 and response.content:
                        target_path.write_bytes(response.content)
                        local_path = target_path
                        resolved_url = export_url
                        status = "downloaded"
                    else:
                        status = "broken_or_restricted"
                        note = "spreadsheet_export_failed"
            elif any(lowered.endswith(ext) for ext in DIRECT_FILE_EXTENSIONS):
                suffix = Path(urlparse(src).path).suffix or ".bin"
                target_path = RAW_EMBED_DIR / f"{doc_slug}-{kind}-{index}{suffix}"
                response = self.session.get(src, timeout=DEFAULT_TIMEOUT)
                response.raise_for_status()
                target_path.write_bytes(response.content)
                local_path = target_path
                status = "downloaded"
        except Exception as exc:  # noqa: BLE001
            note = str(exc)
            status = "error"

        return EmbeddedResource(
            kind=kind,
            source_url=src,
            resolved_url=resolved_url,
            local_path=str(local_path) if local_path and local_path.exists() else None,
            mime_type=mimetypes.guess_type(str(local_path))[0] if local_path else None,
            status=status,
            note=note,
        )

    def _extract_iframes(self, container, doc_slug: str) -> list[EmbeddedResource]:
        resources: list[EmbeddedResource] = []
        for index, frame in enumerate(container.select("iframe[src]"), start=1):
            src = frame.get("src", "").strip()
            if not src:
                continue
            abs_src = absolutize(src) if src.startswith("/") else src
            resources.append(self._download_embedded_resource(abs_src, doc_slug, index, "iframe"))
        return resources

    def _extract_file_links(self, container, doc_slug: str) -> list[EmbeddedResource]:
        resources: list[EmbeddedResource] = []
        seen: set[str] = set()
        index = 0
        for anchor in container.select("a[href]"):
            href = anchor.get("href", "").strip()
            if not href:
                continue
            abs_href = absolutize(href) if href.startswith("/") else href
            lowered = abs_href.lower()
            if lowered in seen:
                continue
            if not (
                any(lowered.endswith(ext) for ext in DIRECT_FILE_EXTENSIONS)
                or "google.com/file/d/" in lowered
                or "docs.google.com/document/d/" in lowered
                or "docs.google.com/spreadsheets/d/" in lowered
            ):
                continue
            seen.add(lowered)
            index += 1
            resources.append(self._download_embedded_resource(abs_href, doc_slug, index, "file"))
        return resources

    def _dedupe_resources(self, resources: list[EmbeddedResource]) -> list[EmbeddedResource]:
        deduped: list[EmbeddedResource] = []
        seen: set[str] = set()
        for resource in resources:
            if resource.source_url in seen:
                continue
            seen.add(resource.source_url)
            deduped.append(resource)
        return deduped

    def _extract_text_from_docx(self, path: Path) -> tuple[str, list[tuple[list[str], list[dict]]]]:
        doc = DocxDocument(str(path))
        paragraphs = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
        tables: list[tuple[list[str], list[dict]]] = []
        for table in doc.tables:
            raw_rows = []
            for row in table.rows:
                values = [normalize_text(cell.text) for cell in row.cells]
                if any(values):
                    raw_rows.append(values)
            if not raw_rows:
                continue
            header = raw_rows[0]
            table_rows: list[dict] = []
            for row in raw_rows[1:]:
                item = {}
                for idx, value in enumerate(row):
                    key = header[idx] if idx < len(header) and header[idx] else f"col_{idx + 1}"
                    item[key] = value
                table_rows.append(item)
            tables.append((header, table_rows))
        return "\n".join(paragraphs), tables

    def _extract_content_from_downloaded_embeds(
        self,
        resources: list[EmbeddedResource],
        doc_id: str,
        source_url: str,
    ) -> tuple[str, list[ExtractedTable], list[str]]:
        extracted_texts: list[str] = []
        tables: list[ExtractedTable] = []
        flags: list[str] = []

        for index, resource in enumerate(resources, start=1):
            if not resource.local_path:
                continue
            path = Path(resource.local_path)
            if not path.exists():
                continue
            suffix = path.suffix.lower()
            try:
                signature = path.read_bytes()[:8]
            except OSError as exc:
                flags.append(f"embed_read_error:{exc}")
                continue

            if signature.startswith(b"%PDF") or suffix == ".pdf":
                flags.append("pdf_embed_extracted")
                try:
                    reader = PdfReader(str(path))
                    pdf_text = []
                    for page in reader.pages[:80]:
                        pdf_text.append(page.extract_text() or "")
                    merged_text = drop_boilerplate_lines(normalize_text("\n".join(pdf_text)))
                    if merged_text:
                        extracted_texts.append(merged_text)
                except Exception as exc:  # noqa: BLE001
                    flags.append(f"pdf_text_error:{exc}")

                try:
                    with pdfplumber.open(str(path)) as pdf:
                        for page_number, page in enumerate(pdf.pages[:40], start=1):
                            for table_index, raw_table in enumerate(page.extract_tables(), start=1):
                                if not raw_table:
                                    continue
                                rows = [row for row in raw_table if row and any(cell for cell in row)]
                                if not rows:
                                    continue
                                header = [normalize_text(str(cell or "")) for cell in rows[0]]
                                normalized_rows = []
                                for row in rows[1:]:
                                    row_values = [
                                        normalize_text(str(row[col_idx])) if col_idx < len(row) and row[col_idx] is not None else ""
                                        for col_idx in range(len(header))
                                    ]
                                    non_empty = [value for value in row_values if value]
                                    if non_empty and non_empty == [str(pos + 1) for pos in range(len(non_empty))]:
                                        continue
                                    item = {}
                                    for col_idx, col_name in enumerate(header):
                                        key = col_name or f"col_{col_idx + 1}"
                                        item[key] = row_values[col_idx] if col_idx < len(row_values) else ""
                                    normalized_rows.append(item)
                                tables.append(
                                    ExtractedTable(
                                        table_id=f"{doc_id}-pdf-table-{index}-{page_number}-{table_index}",
                                        doc_id=doc_id,
                                        title=f"PDF table page {page_number}",
                                        columns=header,
                                        rows=normalized_rows,
                                        source_url=source_url,
                                        extraction_method="pdfplumber",
                                    )
                                )
                except Exception as exc:  # noqa: BLE001
                    flags.append(f"pdf_table_error:{exc}")
            elif suffix in {".xlsx", ".xls"}:
                flags.append("spreadsheet_embed_extracted")
                try:
                    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
                    for sheet in workbook.worksheets:
                        rows = []
                        for row in sheet.iter_rows(values_only=True):
                            values = [normalize_text(str(cell)) if cell is not None else "" for cell in row]
                            if any(values):
                                rows.append(values)
                        if not rows:
                            continue
                        header = rows[0]
                        table_rows = []
                        for row in rows[1:]:
                            item = {}
                            for col_idx, value in enumerate(row):
                                key = header[col_idx] if col_idx < len(header) and header[col_idx] else f"col_{col_idx + 1}"
                                item[key] = value
                            table_rows.append(item)
                        tables.append(
                            ExtractedTable(
                                table_id=f"{doc_id}-xlsx-table-{index}-{slug_from_url(sheet.title)}",
                                doc_id=doc_id,
                                title=sheet.title,
                                columns=header,
                                rows=table_rows,
                                source_url=source_url,
                                extraction_method="openpyxl",
                            )
                        )
                        preview_lines = [sheet.title]
                        for row in rows[:20]:
                            preview_lines.append(" | ".join(value for value in row if value))
                        extracted_texts.append(drop_boilerplate_lines(normalize_text("\n".join(preview_lines))))
                except Exception as exc:  # noqa: BLE001
                    flags.append(f"xlsx_error:{exc}")
            elif suffix == ".docx":
                flags.append("docx_embed_extracted")
                try:
                    docx_text, docx_tables = self._extract_text_from_docx(path)
                    if docx_text:
                        extracted_texts.append(drop_boilerplate_lines(normalize_text(docx_text)))
                    for table_index, (header, rows) in enumerate(docx_tables, start=1):
                        tables.append(
                            ExtractedTable(
                                table_id=f"{doc_id}-docx-table-{index}-{table_index}",
                                doc_id=doc_id,
                                title=f"DOCX table {table_index}",
                                columns=header,
                                rows=rows,
                                source_url=source_url,
                                extraction_method="python-docx",
                            )
                        )
                except Exception as exc:  # noqa: BLE001
                    flags.append(f"docx_error:{exc}")
            else:
                for encoding in ("utf-8", "utf-8-sig", "utf-16"):
                    try:
                        raw_text = path.read_text(encoding=encoding)
                    except (UnicodeDecodeError, OSError):
                        continue
                    cleaned = drop_boilerplate_lines(normalize_text(raw_text))
                    if cleaned:
                        extracted_texts.append(cleaned)
                        flags.append("text_embed_extracted")
                    break

        return "\n\n".join(text for text in extracted_texts if text), tables, flags

    def _build_ocr_variants(self, image_path: str) -> list[tuple[str, object]]:
        try:
            image_bytes = np.fromfile(image_path, dtype=np.uint8)
        except OSError:
            return []
        image = cv2.imdecode(image_bytes, cv2.IMREAD_COLOR)
        if image is None:
            return []

        variants: list[tuple[str, object]] = [("original", image)]
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        variants.append(("gray", gray))
        up2 = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
        up4 = cv2.resize(gray, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        variants.append(("up2_gray", up2))
        variants.append(("up4_gray", up4))

        sharpen_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
        sharpened = cv2.GaussianBlur(gray, (0, 0), 3)
        sharpened = cv2.addWeighted(gray, 1.8, sharpened, -0.8, 0)
        variants.append(("sharpened", sharpened))
        variants.append(("up4_sharpened", cv2.addWeighted(up4, 1.8, cv2.GaussianBlur(up4, (0, 0), 3), -0.8, 0)))

        otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        variants.append(("otsu", otsu))
        variants.append(("up4_otsu", cv2.threshold(up4, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]))

        adaptive = cv2.adaptiveThreshold(
            cv2.GaussianBlur(gray, (3, 3), 0),
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        variants.append(("adaptive", adaptive))
        variants.append(
            (
                "up4_adaptive",
                cv2.adaptiveThreshold(
                    cv2.GaussianBlur(up4, (5, 5), 0),
                    255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    41,
                    15,
                ),
            )
        )
        variants.append(("up4_median", cv2.medianBlur(up4, 3)))
        return variants

    def _ocr_language_score(self, text: str) -> float:
        folded = ascii_fold(text)
        tokens = re.findall(r"[a-z0-9]+", folded)
        if not tokens:
            return 0.0
        score = len(tokens) * 0.5
        score += sum(1.5 for token in tokens if token in {"buoc", "benh", "nhan", "kham", "bhyt", "thuoc", "vien", "phi"})
        score -= sum(0.5 for token in tokens if len(token) == 1 and token not in {"a", "b", "i", "v"})
        return score

    def _normalize_ocr_line(self, line: str) -> str:
        cleaned = normalize_text(line)
        cleaned = re.sub(r"\s+", " ", cleaned)
        cleaned = re.sub(r"[|`~]+", "", cleaned)
        cleaned = cleaned.replace("'", "")
        cleaned = re.sub(r"(?<=\D)(\d)(?=\D|$)", r" \1", cleaned)
        cleaned = re.sub(r"(?<=\bBUOC)\s*(\d+)\b", r" \1", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\b[6G]U0?C\b", "BUOC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bB[AO]?0?C\b", "BUOC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\b[EI]?R?U?OC\b", "BUOC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\b(?:[HN6G]|DU)?R?[0O]?[C]\s*(\d)\b", r"BUOC \1", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bDUOC\s*(\d)\b", r"BUOC \1", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bBUOC\s*([1-9])\b", r"BUOC \1", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bS0\b", "SO", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bTV\b", "TU", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bTHI[' ]?C\b", "THUC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bTHI[' ]?EN\b", "THUC HIEN", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bQ?U[I1]T?R[I1]NH\b", "QUY TRINH", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bKHAMSUC\b", "KHAM SUC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bNOITRU\b", "NOI TRU", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bDANG K[I1]\b", "DANG KI", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bTUQNG\b", "TUONG", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bGIO1\b", "GIOI", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bSOTHU\b", "SO THU", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bCA[NM] LAM SANG\b", "CAN LAM SANG", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bL[I1]NH THUOC\b", "LINH THUOC", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bPHONG KHAM NOP TIEN\b", "PHONG KHAM, NOP TIEN", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s{2,}", " ", cleaned)
        cleaned = cleaned.strip(" ,.-")
        return cleaned

    def _match_process_phrase(self, line: str) -> str:
        folded_line = ascii_fold(line)
        if not folded_line:
            return line
        compact_line = re.sub(r"[^a-z0-9]+", "", folded_line)

        best_phrase = line
        best_score = 0.0
        for phrase in PROCESS_OCR_PHRASES:
            candidate = ascii_fold(phrase)
            compact_candidate = re.sub(r"[^a-z0-9]+", "", candidate)
            score = max(
                SequenceMatcher(None, folded_line, candidate).ratio(),
                SequenceMatcher(None, compact_line, compact_candidate).ratio(),
            )
            if candidate in folded_line or folded_line in candidate:
                score += 0.2
            if compact_candidate in compact_line or compact_line in compact_candidate:
                score += 0.25
            if score > best_score:
                best_score = score
                best_phrase = phrase

        return best_phrase if best_score >= 0.58 else line

    def _refine_ocr_lines(self, lines: list[str], topic_group: str, title: str) -> list[str]:
        refined: list[str] = []
        seen: set[str] = set()

        if topic_group == "process":
            title_match = self._match_process_phrase(title)
            if title_match not in seen:
                refined.append(title_match)
                seen.add(title_match)

        for line in lines:
            normalized = self._normalize_ocr_line(line)
            if not normalized:
                continue
            if topic_group == "process" and len(re.sub(r"[^A-Za-z0-9À-ỹà-ỹ]", "", normalized)) < 4:
                continue
            if topic_group == "process":
                normalized = self._match_process_phrase(normalized)
                folded_normalized = ascii_fold(normalized)
                if folded_normalized in {"buoc", "buoc+", "buoc-"}:
                    continue
                if re.fullmatch(r"[A-Z]{2,6}\s*\d*", folded_normalized.upper()):
                    continue
            if normalized in seen:
                continue
            seen.add(normalized)
            refined.append(normalized)
        return refined

    def _ocr_assets(
        self,
        assets: Iterable[EmbeddedResource],
        topic_group: str = "",
        title: str = "",
    ) -> tuple[list[str], list[str]]:
        facts: list[str] = []
        flags: list[str] = []
        for asset in assets:
            if asset.kind != "image" or not asset.local_path:
                continue
            best_lines: list[str] = []
            best_score = float("-inf")
            best_variant = None
            variant_errors: list[str] = []

            for variant_name, variant in self._build_ocr_variants(asset.local_path):
                try:
                    result, _ = self.ocr_engine(variant)
                except Exception as exc:  # noqa: BLE001
                    variant_errors.append(f"{variant_name}:{exc}")
                    continue
                if not result:
                    continue

                ocr_lines = [normalize_text(item[1]) for item in result if item and len(item) > 1]
                ocr_lines = [line for line in ocr_lines if line]
                if not ocr_lines:
                    continue

                confidence_sum = 0.0
                for item in result:
                    if item and len(item) > 2:
                        try:
                            confidence_sum += float(item[2])
                        except (TypeError, ValueError):
                            continue
                score = confidence_sum + self._ocr_language_score(" ".join(ocr_lines))
                if score > best_score:
                    best_score = score
                    best_lines = ocr_lines
                    best_variant = variant_name

            if not best_lines:
                if variant_errors:
                    flags.append(f"ocr_error:{variant_errors[0]}")
                else:
                    flags.append("ocr_empty")
                continue

            refined_lines = self._refine_ocr_lines(best_lines, topic_group=topic_group, title=title)
            if refined_lines:
                facts.extend(refined_lines[:120])
                flags.append("ocr_used")
                if best_variant and best_variant != "original":
                    flags.append(f"ocr_variant:{best_variant}")
        return facts, flags

    def _should_run_ocr(
        self,
        page_type: str,
        topic_group: str,
        text: str,
        images: list[EmbeddedResource],
        resources: list[EmbeddedResource],
    ) -> bool:
        if not images:
            return False
        if any(resource.local_path for resource in resources):
            return False
        if topic_group == "process":
            return True
        if page_type in {"page", "article"} and len(text) < 500:
            return True
        return False

    def _summarize(self, title: str, text: str) -> str:
        lines = [line for line in text.splitlines() if line]
        preview = " ".join(lines[:3]).strip() or title
        return preview[:500]

    def _normalize_faq_text(self, text: str) -> tuple[str, str]:
        question = ""
        answer = ""
        q_match = re.search(r"Cau hoi:\s*(.*?)\s*Tra loi:", text, re.DOTALL | re.IGNORECASE)
        if q_match:
            question = normalize_text(q_match.group(1))
        a_match = re.search(
            r"Tra loi:\s*(.*?)(?:Nguon|Share:|Cau hoi tu van moi|$)",
            text,
            re.DOTALL | re.IGNORECASE,
        )
        if a_match:
            answer = normalize_text(a_match.group(1))
        combined = []
        if question:
            combined.append("Cau hoi:\n" + question)
        if answer:
            combined.append("Tra loi:\n" + answer)
        cleaned = "\n".join(combined).strip() or text
        return cleaned, question[:120] if question else "FAQ"

    def _extract_contact_points(self, text: str) -> list[str]:
        lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
        contacts: list[str] = []
        contact_keywords = ("hotline", "dien thoai", "lien he", "so dien thoai")
        address_keywords = ("dia chi",)

        def add(value: str) -> None:
            if value and value not in contacts:
                contacts.append(value)

        def find_phones(value: str) -> list[str]:
            results: list[str] = []
            for match in re.finditer(r"(?:\+84|0)[0-9.\-\s()]{8,18}\d", value):
                raw = match.group(0)
                digits = re.sub(r"\D", "", raw)
                if digits.startswith("84") and len(digits) == 11:
                    digits = "0" + digits[2:]
                if 10 <= len(digits) <= 11:
                    results.append(digits)
            return results

        for index, line in enumerate(lines):
            lowered = ascii_fold(line)
            previous = ascii_fold(lines[index - 1]) if index > 0 else ""
            next_line = lines[index + 1] if index + 1 < len(lines) else ""
            search_space = "\n".join(part for part in [line, next_line] if part)

            for email in re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", search_space):
                add(f"email:{email.lower()}")

            if any(keyword in lowered for keyword in address_keywords):
                if ":" in line:
                    address = normalize_text(line.split(":", 1)[1])
                    if address:
                        add(f"address:{address}")
                elif next_line:
                    add(f"address:{next_line}")

            if any(keyword in lowered for keyword in contact_keywords):
                for phone in find_phones(search_space):
                    add(f"phone:{phone}")
            elif any(keyword in previous for keyword in contact_keywords):
                for phone in find_phones(line):
                    add(f"phone:{phone}")
            elif re.fullmatch(r"[\d().+\-\s]{10,20}", line):
                for phone in find_phones(line):
                    add(f"phone:{phone}")

        return contacts

    def _extract_effective_date(self, text: str) -> str | None:
        folded = ascii_fold(text)
        long_date = re.search(r"ngay\s+(\d{1,2})\s+thang\s+(\d{1,2})\s+nam\s*(\d{4})", folded, re.IGNORECASE)
        if long_date:
            day, month, year = (int(long_date.group(1)), int(long_date.group(2)), int(long_date.group(3)))
            return f"{year:04d}-{month:02d}-{day:02d}"
        short_date = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", text)
        if short_date:
            day, month, year = (int(short_date.group(1)), int(short_date.group(2)), int(short_date.group(3)))
            return f"{year:04d}-{month:02d}-{day:02d}"
        return None

    def parse_source(self, source: SourceRecord) -> tuple[KnowledgeDocument, list[ExtractedTable]]:
        html = self.fetch_html(source.source_url)
        soup = BeautifulSoup(html, "lxml")
        container = self.select_main_container(soup, source.page_type)
        title_node = soup.select_one("h1, h2.page-title, h2, h3.entry-title")
        title = normalize_text(title_node.get_text(" ", strip=True)) if title_node else (source.title or source.source_url)
        doc_slug = slug_from_url(source.source_url)
        doc_id = make_id(source.page_type, source.source_url)
        (RAW_HTML_DIR / f"{doc_slug}.html").write_text(html, encoding="utf-8")

        images = self._extract_images(container, doc_slug)
        iframe_root = container if source.page_type in {"page", "article", "faq_detail", "document_listing"} else soup
        iframe_resources = self._extract_iframes(iframe_root, doc_slug)
        file_resources = self._extract_file_links(container, doc_slug)
        resources = self._dedupe_resources([*iframe_resources, *file_resources])

        if source.page_type == "faq_detail":
            text, faq_title = self._extract_structured_faq(container)
            if faq_title:
                title = faq_title
        else:
            text = self._extract_text(container)

        embed_text, tables, embed_flags = self._extract_content_from_downloaded_embeds(resources, doc_id, source.source_url)
        if embed_text:
            text = normalize_text("\n\n".join(part for part in [text, embed_text] if part))

        facts_from_ocr: list[str] = []
        ocr_flags: list[str] = []
        if self._should_run_ocr(source.page_type, source.topic_group, text, images, resources):
            facts_from_ocr, ocr_flags = self._ocr_assets(images, topic_group=source.topic_group, title=title)
        if facts_from_ocr and (source.topic_group == "process" or len(text.strip()) < 120):
            text = normalize_text("\n".join([text, "OCR:\n" + "\n".join(facts_from_ocr)]))

        if source.page_type == "faq_detail":
            text, faq_title = self._normalize_faq_text(text)
            if faq_title:
                title = faq_title

        breadcrumbs = self._extract_breadcrumbs(soup)
        contact_points = (
            self._extract_contact_points(text)
            if source.page_type in {"contact", "department", "department_contact", "faq_detail", "service_page", "form_page"}
            or any(token in ascii_fold(text) for token in ("hotline", "dien thoai", "email", "lien he"))
            else []
        )
        quality_flags = list(dict.fromkeys([*ocr_flags, *embed_flags]))
        if any(resource.status.startswith("broken") or resource.status == "error" for resource in resources):
            quality_flags.append("embed_issue")
        if source.page_type == "faq_detail":
            quality_flags.append("pii_redacted_expected")

        document = KnowledgeDocument(
            doc_id=doc_id,
            source_url=source.source_url,
            page_type=source.page_type,
            topic_group=source.topic_group,
            title=title,
            canonical_title=title.title(),
            breadcrumbs=breadcrumbs,
            text=text,
            summary=self._summarize(title, text),
            department=title if source.page_type in {"department", "department_contact"} else None,
            effective_date=self._extract_effective_date(text),
            contact_points=contact_points,
            assets=[*images, *resources],
            facts=facts_from_ocr[:30],
            tables=[table.table_id for table in tables],
            quality_flags=quality_flags,
            source_status=source.crawl_status,
        )
        return document, tables

    def write_outputs(self, documents: list[KnowledgeDocument], tables: list[ExtractedTable]) -> None:
        write_jsonl(KNOWLEDGE_BASE_DIR / "documents.jsonl", [document.model_dump() for document in documents])
        write_jsonl(KNOWLEDGE_BASE_DIR / "tables.jsonl", [table.model_dump() for table in tables])

        asset_records = []
        for document in documents:
            for asset in document.assets:
                record = asset.model_dump()
                record["doc_id"] = document.doc_id
                asset_records.append(record)

            canonical_path = CANONICAL_DOC_DIR / f"{document.doc_id}.md"
            markdown = [
                f"# {document.canonical_title}",
                "",
                f"- Source: {document.source_url}",
                f"- Page type: {document.page_type}",
                f"- Topic group: {document.topic_group}",
            ]
            if document.department:
                markdown.append(f"- Department: {document.department}")
            if document.effective_date:
                markdown.append(f"- Effective date: {document.effective_date}")
            if document.contact_points:
                markdown.append(f"- Contact points: {', '.join(document.contact_points)}")
            if document.quality_flags:
                markdown.append(f"- Quality flags: {', '.join(document.quality_flags)}")
            markdown.extend(["", "## Summary", "", document.summary, "", "## Text", "", document.text or "(empty)"])
            canonical_path.write_text("\n".join(markdown), encoding="utf-8")

        write_jsonl(KNOWLEDGE_BASE_DIR / "assets.jsonl", asset_records)

    def _reset_outputs(self) -> None:
        for path in [
            CATALOG_DIR / "discovered_sources.jsonl",
            CATALOG_DIR / "katana_discovery.jsonl",
            CATALOG_DIR / "crawl_report.json",
            CATALOG_DIR / "last_run_report.json",
            CATALOG_DIR / "parse_errors.json",
            KNOWLEDGE_BASE_DIR / "documents.jsonl",
            KNOWLEDGE_BASE_DIR / "tables.jsonl",
            KNOWLEDGE_BASE_DIR / "assets.jsonl",
        ]:
            if path.exists():
                path.unlink()
        for path in CANONICAL_DOC_DIR.glob("*.md"):
            path.unlink()

    def run(self, mode: str = "smoke", katana_depth: int | None = None, crawl_duration: str | None = None) -> PipelineReport:
        self._reset_outputs()
        depth = katana_depth or KATANA_DEPTH_BY_MODE.get(mode, 2)
        duration = crawl_duration or KATANA_DURATION_BY_MODE.get(mode, "1m")
        discovered = self.discover_sources(mode=mode, depth=depth, crawl_duration=duration)

        if mode == "smoke":
            discovered_map = {record.source_url: record for record in discovered}
            chosen_sources = [
                discovered_map.get(
                    url,
                    SourceRecord(
                        source_url=url,
                        page_type=self.classify_page_type(url),
                        topic_group=self.classify_topic_group(url),
                        crawl_status="ok",
                    ),
                )
                for url in SMOKE_TARGET_URLS
            ]
        else:
            chosen_sources = [source for source in discovered if source.crawl_status == "ok"]

        documents: list[KnowledgeDocument] = []
        tables: list[ExtractedTable] = []
        errors: list[dict] = []

        for source in chosen_sources:
            try:
                document, table_items = self.parse_source(source)
            except Exception as exc:  # noqa: BLE001
                errors.append({"source_url": source.source_url, "error": str(exc)})
                continue
            documents.append(document)
            tables.extend(table_items)

        self.write_outputs(documents, tables)
        report = PipelineReport(
            mode=mode,
            katana_depth=depth,
            crawl_duration=duration,
            sources_discovered=len(discovered),
            sources_selected=len(chosen_sources),
            documents_written=len(documents),
            tables_written=len(tables),
            errors=len(errors),
        )
        write_json(CATALOG_DIR / "last_run_report.json", report.model_dump())
        write_json(CATALOG_DIR / "parse_errors.json", errors)
        return report
